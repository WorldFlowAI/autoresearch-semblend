"""Generate a professional Microsoft Word patent application document.

Patent formatting best practices:
- 1-inch margins all around
- 14pt bold centered title
- 12pt Times New Roman body
- Double-spaced body text
- Numbered paragraphs throughout specification
- Claims in single-spaced format with proper indentation
- Page numbers centered at bottom
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def create_patent_document(output_path: str) -> None:
    doc = Document()

    # -- Page setup: Letter, 1" margins --
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 2.0

    # -- Heading styles --
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(24 if level == 1 else 18)
        h.paragraph_format.space_after = Pt(12)
        h.paragraph_format.line_spacing = 2.0
        if level == 1:
            h.font.size = Pt(14)
            h.font.bold = True
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            h.font.size = Pt(12)
            h.font.bold = True
            h.font.underline = True
        else:
            h.font.size = Pt(12)
            h.font.bold = True

    # Paragraph counter for numbered paragraphs
    para_num = [0]

    def add_title_page():
        """Create the patent application title/cover page."""
        # Blank lines for spacing
        for _ in range(4):
            doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("UNITED STATES PATENT APPLICATION")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"

        doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(
            "SYSTEM AND METHOD FOR SEMANTIC KEY-VALUE\n"
            "CACHE ROUTING IN LANGUAGE MODEL INFERENCE"
        )
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

        doc.add_paragraph("")
        doc.add_paragraph("")

        # Metadata block
        meta_items = [
            ("Applicant:", "WorldFlow AI, Inc."),
            ("Inventors:", "Zachary Bennett, Coby Litvinsky"),
            ("Date:", "March 2026"),
            ("Filing Type:", "Continuation-in-Part of Filing 1 and Filing 3"),
            ("Version:", "5.0"),
        ]
        for label, value in meta_items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(label + " ")
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
            run = p.add_run(value)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

        doc.add_paragraph("")
        doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Prepared for: David Connaughton, Lambert Patent Law")
        run.italic = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONFIDENTIAL — ATTORNEY-CLIENT PRIVILEGED")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(180, 0, 0)

        doc.add_page_break()

    def numbered_para(text: str, bold_first_sentence: bool = False):
        """Add a numbered paragraph in patent specification style."""
        para_num[0] += 1
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.first_line_indent = Inches(0.5)

        bracket = f"[{para_num[0]:04d}] "
        run = p.add_run(bracket)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        if bold_first_sentence and ". " in text:
            first_sent, rest = text.split(". ", 1)
            run = p.add_run(first_sent + ". ")
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run = p.add_run(rest)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        else:
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        return p

    def section_heading(text: str, level: int = 1):
        """Add a section heading."""
        doc.add_heading(text, level=level)

    def add_claim(claim_num: str, text: str, dependent: bool = False):
        """Add a patent claim with proper formatting."""
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)
        if dependent:
            p.paragraph_format.left_indent = Inches(0.5)

        run = p.add_run(f"{claim_num}. ")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    def add_claim_element(text: str, indent_level: int = 1):
        """Add a claim sub-element with proper indentation."""
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.5 * indent_level)
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # ================================================================
    #  BUILD THE DOCUMENT
    # ================================================================

    add_title_page()

    # --- CROSS-REFERENCE TO RELATED APPLICATIONS ---
    section_heading("CROSS-REFERENCE TO RELATED APPLICATIONS")

    numbered_para(
        'This application is a continuation-in-part of U.S. Patent Application '
        'Serial No. [TO BE ASSIGNED], titled "Multi-Model Semantic Caching '
        'Architecture for AI Applications" (Filing 1), originally filed under '
        'COTI LLC, assignment transfer to WorldFlow AI, Inc. pending. Priority '
        'is claimed under 35 U.S.C. \u00a7 120 for all commonly disclosed subject '
        'matter, including the multi-tier semantic cache hierarchy, semantic '
        'similarity scoring for AI inference query deduplication, and the '
        'configurable similarity threshold mechanism.'
    )

    numbered_para(
        'This application is related to U.S. Patent Application Serial No. '
        '[TO BE ASSIGNED], titled "Personalization Contamination Detection '
        'System for Cross-User Cache Management" (Filing 2), filed '
        'December 19, 2025. No priority is claimed from Filing 2.'
    )

    numbered_para(
        'This application is a continuation-in-part of U.S. Patent Application '
        'Serial No. [TO BE ASSIGNED], titled "Semantic Compositional Caching '
        'System for AI-Powered Visual Content Generation" (Filing 3), filed '
        'February 4, 2026. Priority is claimed under 35 U.S.C. \u00a7 120 for '
        'commonly disclosed subject matter including the dual-GPU execution '
        'architecture, three-tier memory hierarchy, and GPU-resident vector indexing.'
    )

    numbered_para(
        'New matter introduced herein: (a) semantic fleet router that dispatches '
        'inference queries to GPU workers based on semantic similarity between the '
        'query and each worker\u2019s cached key-value tensor states, with event-driven '
        'cache state synchronization; (b) cascade discovery architecture combining '
        'exact hash lookup with semantic approximate nearest neighbor search in a '
        'strictly additive pipeline; (c) multi-region prompt sampling for '
        'representative embedding of long-context inputs; (d) position-invariant '
        'text representation for reorder-tolerant cache matching; (e) tenant '
        'affinity scoring in semantic cache routing; (f) architecture-specific '
        'layer divergence profiling as applied to per-model selective recomputation; '
        '(g) end-of-sequence collapse prevention via minimum fresh token guarantee.'
    )

    numbered_para(
        'CRITICAL NOTE FOR COUNSEL: The COTI LLC to WorldFlow AI assignment '
        'transfer for Filing 1 must be recorded at the USPTO before this CIP '
        'establishes its priority chain. This is a blocking item.'
    )

    # --- NOTE ON PRIOR ART ---
    section_heading("NOTE ON PRIOR ART INCORPORATED BY REFERENCE")

    numbered_para(
        'The system described herein incorporates several known techniques as '
        'building blocks within the preferred embodiment: rotary position embedding '
        'correction for non-contiguous key-value reuse (as described in MEPIC, '
        'arxiv 2512.16822, December 2025; EPIC, arxiv 2410.15332, October 2024); '
        'selective layer recomputation for cached key-value fusion (as described in '
        'CacheBlend, EuroSys 2025); token-level edit alignment for donor-query '
        'matching (as described in KVShare, arxiv 2503.16525, March 2025); and '
        'semantic similarity matching for cache lookup (as described in GPTCache, '
        'NLP-OSS 2023; SemShareKV, arxiv 2509.24832, September 2025). These known '
        'techniques are not independently claimed herein. The present invention\u2019s '
        'novelty resides in the specific combination, system architecture, and '
        'novel methods described in the claims below.'
    )

    # --- ABSTRACT ---
    section_heading("ABSTRACT OF THE DISCLOSURE")

    numbered_para(
        'A computer-implemented system and method for routing artificial intelligence '
        'inference queries across a plurality of GPU workers based on semantic '
        'similarity between incoming queries and cached key-value tensor states '
        'maintained on each worker. A semantic fleet router maintains a fleet-level '
        'vector index synchronized in real time via a distributed event plane, '
        'enabling cache-affinity-aware routing without direct worker polling. A '
        'cascade discovery architecture combines exact hash lookup with semantic '
        'approximate nearest neighbor search in a strictly additive pipeline that '
        'only activates semantic search upon exact-match miss. A multi-region prompt '
        'sampling method generates representative embeddings of long-context inputs '
        'by extracting text from head, middle, and tail regions rather than '
        'prefix-only truncation. A position-invariant text encoding method sorts '
        'constituent segments prior to embedding, enabling cache matching across '
        'semantically equivalent but differently-ordered inputs. Tenant affinity '
        'scoring boosts same-tenant matches in routing decisions based on empirical '
        'observation of higher intra-tenant semantic overlap. The system integrates '
        'with known partial key-value reuse and position correction techniques to '
        'reduce prefill computation on cache hits while preserving output quality.'
    )

    # --- BACKGROUND ---
    section_heading("BACKGROUND OF THE INVENTION")

    section_heading("Technical Field", level=2)

    numbered_para(
        'The present invention relates to computer-implemented systems and methods '
        'for routing artificial intelligence inference queries across GPU-equipped '
        'compute infrastructure based on semantic cache affinity. More specifically, '
        'the invention relates to fleet-level routing of inference queries to GPU '
        'workers based on the semantic content of each worker\u2019s cached key-value '
        'tensor states, and to methods for generating representative embeddings of '
        'long-context inputs for use in semantic cache matching.'
    )

    section_heading("Description of Related Art", level=2)

    numbered_para(
        'Modern AI inference systems face escalating costs driven by increasing '
        'model sizes, longer context windows, and growing query volumes. The '
        'inference prefill phase, which computes attention key-value (KV) tensor '
        'states for the input context, dominates time-to-first-token (TTFT) '
        'latency. For long-context inputs, prefill computation can exceed several '
        'seconds, creating a significant bottleneck for interactive applications.'
    )

    numbered_para(
        'Several approaches exist for reducing redundant prefill computation. '
        'Exact prefix caching systems cache KV tensors and reuse them when an '
        'incoming query shares an identical token prefix with a cached entry. '
        'These systems achieve zero overhead on hits but have limited coverage: '
        'empirical analysis of large-scale conversational datasets shows that only '
        'a small fraction of consecutive queries share near-exact prefixes.'
    )

    numbered_para(
        'Application-level semantic caches match complete prompts by semantic '
        'similarity and return cached responses, but have no visibility into KV '
        'tensor states and cannot enable partial reuse of attention computations.'
    )

    numbered_para(
        'Semantic KV cache sharing systems enable reuse of cached KV tensors '
        'between semantically similar prompts on a single inference instance, '
        'using techniques such as token-level matching and selective recomputation. '
        'Position-independent caching systems decouple cached KV from absolute '
        'positions using rotary position embedding correction. Other systems '
        'enable fusion of pre-computed KV caches via selective layer recomputation.'
    )

    numbered_para(
        'However, no prior system addresses the fleet-level routing problem: '
        'given a cluster of GPU workers each maintaining local KV caches, '
        'determining which worker to route an incoming query to based on the '
        'semantic content of that worker\u2019s cached states. Prior KV-aware routers '
        'match on exact token prefixes only. Prior semantic routers route between '
        'different models, not between instances of the same model based on KV '
        'cache content.'
    )

    numbered_para(
        'Additionally, no prior system addresses the embedding quality problem '
        'for long-context semantic cache matching. Existing embedding approaches '
        'truncate long inputs to a fixed-length prefix, which for long-context '
        'inputs captures only a small fraction of the content and produces '
        'non-representative embeddings that fail to match semantically similar '
        'cached entries.'
    )

    # --- SUMMARY ---
    section_heading("SUMMARY OF THE INVENTION")

    numbered_para(
        'In accordance with the present invention, a system and method are '
        'provided for fleet-level semantic routing of inference queries based on '
        'KV cache affinity, cascade discovery combining exact and semantic '
        'matching, multi-region prompt sampling for long-context embedding, and '
        'position-invariant text encoding for reorder-tolerant matching. The system '
        'integrates with known partial KV reuse and position correction techniques '
        'to achieve significant TTFT speedup on cache hits while preserving output '
        'quality. The system is strictly additive to existing exact-prefix '
        'routing\u2014it activates only when exact matching fails, never replacing or '
        'degrading existing cache paths.'
    )

    # --- DETAILED DESCRIPTION ---
    section_heading("DETAILED DESCRIPTION OF PREFERRED EMBODIMENTS")

    section_heading("4.1 System Architecture Overview (FIG. 50)", level=2)

    numbered_para(
        'Referring now to FIG. 50, a fleet-level semantic inference routing system '
        'is illustrated. The system comprises a plurality of GPU workers, each '
        'maintaining a local key-value tensor cache, and a semantic fleet router '
        'that routes incoming inference queries to workers based on cache affinity. '
        'In the preferred embodiment, the system is deployed on GPU infrastructure '
        'within a container orchestration platform with a plurality of GPU worker '
        'pods, a fleet router pod, and a distributed message bus for real-time '
        'cache state synchronization.'
    )

    numbered_para(
        'Each GPU worker maintains a local vector similarity index over its cached '
        'entries. In the preferred embodiment, the index auto-selects backend by '
        'entry count: exact search for small populations, graph-based approximate '
        'nearest neighbor search for large populations, with CPU fallback when GPU '
        'acceleration is unavailable. Cache entries comprise cached KV tensor states '
        'indexed by dense vector embeddings generated by a lightweight embedding model.'
    )

    section_heading("4.2 Cascade Discovery Engine (FIG. 51)", level=2)

    numbered_para(
        'Referring now to FIG. 51, the cascade discovery engine implements a '
        'strictly additive two-stage pipeline. Stage 1 performs exact chunk-hash '
        'lookup: the token-level cache computes hashes over fixed-size token chunks '
        'and compares against cached chunk identifiers. When Stage 1 finds a match, '
        'the cached result is returned immediately without any embedding '
        'computation\u2014zero overhead. When Stage 1 misses, and only then, Stage 2 '
        'activates: the semantic pipeline embeds the prompt using the '
        'position-invariant encoding method described in Section 4.6, then performs '
        'approximate nearest neighbor search against a fleet-level vector index. '
        'Results exceeding a configurable similarity threshold are classified as '
        'semantic cache hits. The semantic stage never replaces or degrades the '
        'exact-match path.'
    )

    numbered_para(
        'This strictly additive architecture is a key differentiator: the system '
        'adds semantic matching capability to any existing exact-match cache system '
        'without modifying or degrading the existing system\u2019s behavior.'
    )

    section_heading("4.3 Semantic Fleet Router (FIG. 52)", level=2)

    numbered_para(
        'Referring now to FIG. 52, the semantic fleet router operates as a '
        'centralized routing plane. On each incoming inference query, the router:'
    )

    numbered_para(
        '(1) Generates a dense vector embedding of the query using the multi-region '
        'sampling method (Section 4.5) and position-invariant encoding (Section 4.6);'
    )

    numbered_para(
        '(2) Searches the fleet-level vector index for cached entries with highest '
        'cosine similarity to the query embedding;'
    )

    numbered_para(
        '(3) Applies tenant affinity scoring (Section 4.7) to boost same-tenant matches;'
    )

    numbered_para(
        '(4) If a match exceeds the configurable similarity threshold, dispatches '
        'the query to the worker holding the matched cached entry, including a donor '
        'hint identifying the specific cached entry;'
    )

    numbered_para(
        '(5) If no match exceeds the threshold, falls back to least-loaded worker '
        'selection (cold fallback).'
    )

    numbered_para(
        'The fleet-level vector index is maintained via event-driven synchronization: '
        'each GPU worker publishes cache registration events (when a new entry is '
        'cached) and cache removal events (when an entry is evicted) to a distributed '
        'message bus. The router\u2019s event consumer processes these events to add or '
        'remove entries from the fleet index in real time. This eliminates the need '
        'for the router to poll individual workers, enabling constant-time routing '
        'decisions regardless of fleet size.'
    )

    numbered_para(
        'In the preferred embodiment, the routing service is implemented as an HTTP '
        'API gateway, the message bus uses a lightweight publish-subscribe protocol, '
        'and the fleet-level vector index uses the same auto-adaptive backend '
        'selection as the per-worker indexes.'
    )

    section_heading("4.4 Partial KV Reuse (Preferred Embodiment)", level=2)

    numbered_para(
        'On semantic cache hit, the selected GPU worker performs partial key-value '
        'tensor reuse using known techniques: the worker\u2019s local semantic pipeline '
        'finds the best-matching local donor, computes token-level alignment between '
        'donor and query, injects aligned donor KV cache for matching tokens, and '
        'selectively recomputes divergent layers. In the preferred embodiment, '
        'rotary position embedding correction is applied to cached attention key '
        'matrices to handle position differences, and architecture-specific layer '
        'recomputation masks are applied to ensure quality preservation. '
        'End-of-sequence collapse is prevented by ensuring a minimum number of '
        'fresh tokens are computed from the current query before any donor KV injection.'
    )

    section_heading("4.5 Multi-Region Prompt Sampling", level=2)

    numbered_para(
        'For long-context inputs where the token count substantially exceeds the '
        'embedding model\u2019s effective window, a multi-region sampling method extracts '
        'text from multiple positions within the input token sequence. In the '
        'preferred embodiment, text is sampled from three regions\u2014the head, middle, '
        'and tail of the input\u2014at configurable proportions (e.g., 40%, 30%, 30%). '
        'This produces a more representative embedding than simple prefix truncation, '
        'which for long inputs may capture only a small fraction of the total content.'
    )

    numbered_para(
        'This is a critical innovation: empirical validation shows that prefix-only '
        'truncation yields near-zero cache hit rates on cross-instruction workloads '
        'at context lengths exceeding a threshold number of tokens, because the '
        'prefix captures only the system prompt and initial context\u2014not the document '
        'body that distinguishes one query from another. Multi-region sampling '
        'captures content from throughout the input, restoring full hit rates and '
        'corresponding speedups.'
    )

    section_heading("4.6 Position-Invariant Text Encoding", level=2)

    numbered_para(
        'Before embedding, the input text is processed through a position-invariant '
        'encoding step: the text is segmented into constituent parts (e.g., '
        'sentences), the parts are sorted alphabetically, and the sorted text is '
        'concatenated to form the embedding input. This produces identical embeddings '
        'regardless of the ordering of segments in the original input.'
    )

    numbered_para(
        'This is important for cache matching in scenarios where the same content '
        'appears in different orderings across queries\u2014for example, when a '
        'retrieval-augmented generation system returns the same documents in '
        'different order for different queries, or when conversational context is '
        'assembled differently across turns.'
    )

    section_heading("4.7 Tenant Affinity Scoring", level=2)

    numbered_para(
        'In multi-tenant deployments, the routing score for each candidate cached '
        'entry includes a configurable tenant affinity bonus added when the query\u2019s '
        'tenant identifier matches the cached entry\u2019s tenant identifier. This '
        'reflects the empirical finding from analysis of large-scale conversational '
        'data: same-user consecutive query pairs exhibit substantially higher semantic '
        'overlap than cross-user pairs. The affinity bonus is tunable per deployment '
        'to balance routing precision against cross-tenant cache utilization.'
    )

    # --- BRIEF DESCRIPTION OF THE DRAWINGS ---
    section_heading("BRIEF DESCRIPTION OF THE DRAWINGS")

    numbered_para(
        'FIG. 50 is a system architecture diagram illustrating the fleet-level '
        'semantic inference routing system with a plurality of GPU workers, a fleet '
        'router, and a distributed event plane.'
    )

    numbered_para(
        'FIG. 51 is a system architecture diagram illustrating the cascade discovery '
        'engine with the two-stage strictly additive pipeline.'
    )

    numbered_para(
        'FIG. 52 is a system architecture diagram illustrating the semantic fleet '
        'router with event-driven cache state synchronization and multi-path routing.'
    )

    # --- CLAIMS ---
    section_heading("CLAIMS")

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("What is claimed is:")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # --- System Claims (8001-8008) ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("System Claims")
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Claim 8001 (independent)
    add_claim(
        "Claim 1",
        "A computer-implemented system for routing artificial intelligence "
        "inference queries across a plurality of GPU workers based on semantic "
        "cache affinity, the system comprising:"
    )
    add_claim_element(
        "(a) a plurality of GPU workers, each GPU worker maintaining a local "
        "cache of key-value tensor states computed during prior inference "
        "operations, and each GPU worker publishing cache state change events "
        "to a distributed message bus upon caching or evicting key-value tensor states;"
    )
    add_claim_element(
        "(b) a semantic fleet router configured to receive inference queries "
        "from client applications and to maintain a fleet-level vector index "
        "of cached entries across all GPU workers, the fleet-level vector index "
        "being updated in real time from cache state change events received "
        "from the distributed message bus; and"
    )
    add_claim_element(
        "(c) a cascade discovery engine within the semantic fleet router "
        "comprising an exact hash lookup stage and a semantic search stage, "
        "the semantic search stage activating only when the exact hash lookup "
        "stage yields no match, and the semantic search stage querying the "
        "fleet-level vector index to identify cached entries semantically "
        "similar to the inference query;"
    )
    add_claim_element(
        "wherein the semantic fleet router dispatches each inference query to "
        "the GPU worker holding the highest-scoring cached entry when the "
        "similarity exceeds a configurable threshold, and dispatches to a "
        "least-loaded GPU worker when no cached entry exceeds the threshold."
    )

    # Claim 8002
    add_claim(
        "Claim 2",
        "The system of claim 1, wherein the fleet-level vector index is "
        "synchronized with the plurality of GPU workers exclusively via the "
        "distributed message bus, such that the semantic fleet router determines "
        "cache affinity for incoming queries without directly polling any GPU worker.",
        dependent=True
    )

    # Claim 8003
    add_claim(
        "Claim 3",
        "The system of claim 1, wherein the exact hash lookup stage of the "
        "cascade discovery engine computes chunk-level hashes of the incoming "
        "inference query using fixed-size token chunks and performs hash "
        "comparison against cached chunk identifiers, and upon identifying an "
        "exact match, bypasses the semantic search stage entirely, incurring "
        "zero embedding computation overhead.",
        dependent=True
    )

    # Claim 8004
    add_claim(
        "Claim 4",
        "The system of claim 1, wherein the semantic search stage generates a "
        "dense vector embedding of the inference query using position-invariant "
        "encoding in which text segments extracted from the query are sorted "
        "prior to encoding, producing identical embeddings regardless of segment "
        "ordering in the query.",
        dependent=True
    )

    # Claim 8005
    add_claim(
        "Claim 5",
        "The system of claim 1, wherein the semantic fleet router computes a "
        "routing score for each candidate cached entry by combining a cosine "
        "similarity score with a tenant affinity bonus added when the query\u2019s "
        "tenant identifier matches the cached entry\u2019s tenant identifier, the "
        "tenant affinity bonus being a configurable value.",
        dependent=True
    )

    # Claim 8006
    add_claim(
        "Claim 6",
        "The system of claim 1, wherein each cache state change event published "
        "by a GPU worker to the distributed message bus comprises a cache "
        "registration event including a dense vector embedding and a worker "
        "identifier, or a cache removal event including a cache entry identifier "
        "and a worker identifier.",
        dependent=True
    )

    # Claim 8007
    add_claim(
        "Claim 7",
        "The system of claim 1, wherein the semantic fleet router, upon "
        "dispatching an inference query to a selected GPU worker, transmits a "
        "donor hint identifying the specific cached entry to the selected GPU "
        "worker, enabling the worker to locate the cached entry without "
        "performing a local search.",
        dependent=True
    )

    # Claim 8008
    add_claim(
        "Claim 8",
        "The system of claim 1, wherein the fleet-level vector index "
        "auto-selects a search backend based on the number of indexed entries, "
        "using exact search for entry counts below a configurable threshold and "
        "graph-based approximate nearest neighbor search for entry counts at or "
        "above the threshold.",
        dependent=True
    )

    # --- Method Claims (8009-8016) ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Method Claims")
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Claim 8009 (independent)
    add_claim(
        "Claim 9",
        "A computer-implemented method for routing inference queries across a "
        "plurality of GPU workers based on semantic similarity to cached "
        "key-value tensor states, the method comprising:"
    )
    add_claim_element(
        "(a) maintaining, at a semantic fleet router, a fleet-level vector index "
        "of cached key-value tensor states across the plurality of GPU workers, "
        "the fleet-level vector index being updated by consuming cache "
        "registration events and cache removal events from a distributed message bus;"
    )
    add_claim_element(
        "(b) upon receiving an inference query, generating a dense vector "
        "embedding of the inference query;"
    )
    add_claim_element(
        "(c) performing a cascade discovery operation comprising first executing "
        "an exact hash lookup against cached query identifiers and, only when "
        "the exact hash lookup yields no match, executing a semantic search "
        "against the fleet-level vector index;"
    )
    add_claim_element(
        "(d) when the semantic search identifies a cached entry with similarity "
        "exceeding a configurable threshold, dispatching the inference query to "
        "the GPU worker holding the cached entry along with a donor hint "
        "identifying the cached entry; and"
    )
    add_claim_element(
        "(e) when no cached entry exceeds the threshold, dispatching the "
        "inference query to a least-loaded GPU worker;"
    )
    add_claim_element(
        "wherein the cascade discovery operation ensures that the semantic "
        "search adds capability without degrading the performance of exact "
        "hash matching."
    )

    # Claim 8010
    add_claim(
        "Claim 10",
        "The method of claim 9, further comprising, upon completion of inference "
        "on a GPU worker, the GPU worker publishing a cache registration event "
        "to the distributed message bus, the event including a dense vector "
        "embedding of the completed query and the worker\u2019s identifier, enabling "
        "the semantic fleet router to route subsequent semantically similar "
        "queries to that worker.",
        dependent=True
    )

    # Claim 8011
    add_claim(
        "Claim 11",
        "The method of claim 9, further comprising, upon eviction of a cached "
        "entry from a GPU worker, the GPU worker publishing a cache removal "
        "event to the distributed message bus, enabling the semantic fleet "
        "router to remove the entry from the fleet-level vector index and avoid "
        "routing queries to stale cached entries.",
        dependent=True
    )

    # Claim 8012
    add_claim(
        "Claim 12",
        "The method of claim 9, wherein generating the dense vector embedding "
        "comprises applying position-invariant encoding in which text segments "
        "of the query are sorted alphabetically prior to encoding by an "
        "embedding model.",
        dependent=True
    )

    # Claim 8013
    add_claim(
        "Claim 13",
        "The method of claim 9, wherein generating the dense vector embedding "
        "further comprises, when the inference query exceeds a predetermined "
        "length, extracting text from a plurality of regions within the query "
        "rather than truncating to a prefix, the regions including at least a "
        "head region, a middle region, and a tail region.",
        dependent=True
    )

    # Claim 8014
    add_claim(
        "Claim 14",
        "The method of claim 9, wherein dispatching the inference query to the "
        "GPU worker holding the cached entry further comprises applying a tenant "
        "affinity bonus to the similarity score when the query\u2019s tenant "
        "identifier matches the cached entry\u2019s tenant identifier, the bonus "
        "being a configurable value reflecting higher intra-tenant semantic "
        "overlap observed in production multi-tenant workloads.",
        dependent=True
    )

    # Claim 8015
    add_claim(
        "Claim 15",
        "The method of claim 9, wherein the cascade discovery operation is "
        "strictly additive to an existing exact-prefix cache system, such that "
        "the exact hash lookup stage operates identically to the existing system "
        "and the semantic search stage only handles queries that the existing "
        "system would have sent to cold fallback.",
        dependent=True
    )

    # Claim 8016
    add_claim(
        "Claim 16",
        "The method of claim 9, wherein the fleet-level vector index "
        "auto-selects a search backend based on entry count, using exact search "
        "for small populations and approximate nearest neighbor search for "
        "large populations.",
        dependent=True
    )

    # --- Multi-Region Sampling and Encoding Claims (8017-8024) ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Multi-Region Sampling and Encoding Claims")
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Claim 8017 (independent)
    add_claim(
        "Claim 17",
        "A computer-implemented method for generating representative dense "
        "vector embeddings of long-context inference queries for use in semantic "
        "cache matching, the method comprising:"
    )
    add_claim_element(
        "(a) receiving an inference query comprising a token sequence exceeding "
        "a predetermined length threshold;"
    )
    add_claim_element(
        "(b) extracting text from a plurality of non-overlapping regions within "
        "the token sequence, the regions including at least a head region "
        "beginning at the start of the sequence, a middle region centered "
        "within the sequence, and a tail region ending at the end of the sequence;"
    )
    add_claim_element(
        "(c) concatenating the extracted text from the plurality of regions to "
        "form a composite text representation;"
    )
    add_claim_element(
        "(d) sorting segments of the composite text representation to produce "
        "a position-invariant encoding; and"
    )
    add_claim_element(
        "(e) generating a dense vector embedding from the position-invariant "
        "encoding using an embedding model;"
    )
    add_claim_element(
        "wherein the multi-region extraction produces a more representative "
        "embedding than prefix-only truncation for long-context inputs, enabling "
        "semantic cache matching across queries that share content distributed "
        "throughout the input."
    )

    # Claim 8018
    add_claim(
        "Claim 18",
        "The method of claim 17, wherein the proportions of text extracted from "
        "each region are configurable per deployment.",
        dependent=True
    )

    # Claim 8019
    add_claim(
        "Claim 19",
        "The method of claim 17, wherein for inference queries below the "
        "predetermined length threshold, text is extracted from the full query "
        "without multi-region sampling, and the position-invariant encoding is "
        "applied to the full extracted text.",
        dependent=True
    )

    # Claim 8020
    add_claim(
        "Claim 20",
        "The method of claim 17, wherein the position-invariant encoding "
        "comprises segmenting the text into sentence-level units, sorting the "
        "units alphabetically, and concatenating the sorted units.",
        dependent=True
    )

    # Claim 8021
    add_claim(
        "Claim 21",
        "The method of claim 17, wherein the dense vector embedding is used by "
        "a semantic fleet router to search a fleet-level vector index for cached "
        "key-value tensor states across a plurality of GPU workers, enabling "
        "routing of the inference query to the GPU worker whose cached state is "
        "most semantically similar to the query.",
        dependent=True
    )

    # Claim 8022
    add_claim(
        "Claim 22",
        "The method of claim 17, wherein the predetermined length threshold is "
        "set based on the effective context window of the embedding model, such "
        "that multi-region sampling activates only when the query length would "
        "cause prefix-only truncation to capture less than a configurable "
        "fraction of the total input content.",
        dependent=True
    )

    # Claim 8023
    add_claim(
        "Claim 23",
        "The system of claim 1, wherein the system is strictly additive to an "
        "existing exact-prefix KV cache routing system, such that the cascade "
        "discovery engine\u2019s exact hash lookup stage corresponds to the existing "
        "system\u2019s matching behavior and the semantic search stage handles only "
        "queries that the existing system would route to cold fallback, ensuring "
        "zero performance regression on workloads served by the existing system.",
        dependent=True
    )

    # Claim 8024
    add_claim(
        "Claim 24",
        "The system of claim 1, wherein each GPU worker further maintains a "
        "local vector similarity index over its own cached entries, and upon "
        "receiving a dispatched query with a donor hint, the GPU worker uses "
        "the donor hint to locate the cached entry and perform partial key-value "
        "tensor reuse without re-searching its local index.",
        dependent=True
    )

    # --- GLOSSARY ---
    section_heading("GLOSSARY OF KEY TERMS")

    glossary_terms = [
        ("Semantic Fleet Router",
         "A routing service that dispatches inference queries to GPU workers "
         "based on semantic similarity between the query and each worker\u2019s "
         "cached key-value tensor states. Maintains a fleet-level vector index "
         "synchronized via distributed event plane."),
        ("Fleet-Level Vector Index",
         "A centralized index of cached entry embeddings across all GPU workers, "
         "updated in real time via cache state change events. Enables "
         "cache-affinity routing without direct worker polling."),
        ("Cascade Discovery Engine",
         "A strictly additive two-stage pipeline: Stage 1 exact chunk-hash "
         "lookup; Stage 2 semantic approximate nearest neighbor search, "
         "activating only on Stage 1 miss. Adds semantic capability without "
         "degrading existing exact-match behavior."),
        ("Cache Affinity Routing",
         "Dispatching an inference query to the GPU worker whose cached "
         "key-value tensor states are most semantically similar to the query, "
         "rather than selecting a worker at random or by load alone."),
        ("Distributed Event Plane",
         "Message bus through which workers publish cache registration and "
         "removal events, enabling the fleet router to maintain the fleet-level "
         "vector index without polling."),
        ("Tenant Affinity Scoring",
         "A configurable bonus added to routing scores for same-tenant matches, "
         "reflecting higher intra-tenant semantic overlap in multi-tenant workloads."),
        ("Multi-Region Prompt Sampling",
         "Method for generating representative embeddings of long-context inputs "
         "by extracting text from head, middle, and tail regions rather than "
         "truncating to a prefix."),
        ("Position-Invariant Encoding",
         "Text representation method that sorts constituent segments prior to "
         "embedding, producing identical vectors regardless of segment ordering."),
        ("Donor Hint",
         "Metadata transmitted from the fleet router to a selected GPU worker "
         "identifying the specific cached entry to reuse, eliminating the need "
         "for the worker to perform a local search."),
        ("Configurable Similarity Threshold",
         "Cosine similarity value above which a cached entry is considered a "
         "semantic cache hit. Tunable per deployment to balance cache hit rate "
         "against match quality."),
    ]

    for term, definition in glossary_terms:
        numbered_para(f"{term}. {definition}", bold_first_sentence=True)

    # --- PRIOR ART DIFFERENTIATION ---
    section_heading("PRIOR ART DIFFERENTIATION")

    section_heading("Prior Art Landscape", level=2)

    numbered_para(
        'The following prior art systems are relevant to this filing:'
    )

    prior_art_items = [
        ("Exact-prefix KV cache routing",
         "These systems route queries to workers based on exact token prefix "
         "matches. They achieve zero overhead on hits but have limited coverage\u2014"
         "empirically, only a small fraction of queries share exact prefixes with "
         "cached entries. The present invention extends coverage to all semantically "
         "similar queries above a configurable threshold while preserving the "
         "existing exact-match behavior unchanged."),
        ("Semantic prompt caching (GPTCache, NLP-OSS 2023)",
         "Caches and returns complete LLM responses based on prompt-level semantic "
         "similarity. Operates at the application layer with no visibility into KV "
         "tensor states. Cannot enable partial reuse of attention computations. The "
         "present invention operates at the KV tensor level, enabling partial "
         "prefill reuse, not response-level caching."),
        ("Semantic KV cache sharing (SemShareKV, 2025; KVShare, 2025)",
         "Enable reuse of cached KV tensors between semantically similar prompts "
         "on a SINGLE inference instance. Do not address fleet-level routing across "
         "multiple workers. The present invention solves the fleet routing problem: "
         "given N workers each with different cached states, route each query to "
         "the best worker."),
        ("Position-independent caching (MEPIC, 2025; EPIC, 2024)",
         "Decouple cached KV from absolute positions via rotary position embedding "
         "correction. These techniques are incorporated as known building blocks in "
         "the preferred embodiment but are not independently claimed."),
        ("Selective layer recomputation (CacheBlend, EuroSys 2025)",
         "Fuses pre-computed KV caches via selective recomputation of a fraction of "
         "tokens per layer. Incorporated as a known building block in the preferred "
         "embodiment but not independently claimed. The present invention extends "
         "this with per-architecture divergence profiling."),
        ("Semantic model routing (vLLM Semantic Router, 2025\u20132026)",
         "Routes queries between DIFFERENT models based on query semantics "
         "(Mixture-of-Models). Does NOT route between instances of the same model "
         "based on KV cache content. The present invention solves a fundamentally "
         "different problem: routing to the best CACHE, not the best MODEL."),
    ]

    for title, desc in prior_art_items:
        numbered_para(f"{title}: {desc}", bold_first_sentence=True)

    section_heading("Key Differentiators", level=2)

    differentiators = [
        ("Fleet-Level Semantic KV Cache Routing (Claims 1\u20132, 5\u20137, 9\u201311, 14)",
         "No prior system routes inference queries across a fleet of GPU workers "
         "based on semantic similarity between the query and each worker\u2019s cached "
         "KV tensor states, with event-driven fleet index synchronization. This "
         "is the core novel contribution."),
        ("Cascade Discovery\u2014Strictly Additive (Claims 1(c), 3, 9(c), 15, 23)",
         "No prior system combines exact chunk-hash lookup with semantic vector "
         "search in a strictly additive pipeline where semantic search activates "
         "only on exact miss. Existing systems are either exact-only or semantic-only."),
        ("Multi-Region Prompt Sampling (Claims 13, 17\u201322)",
         "No prior system addresses the embedding truncation problem for long-context "
         "semantic cache matching by sampling from multiple regions (head, middle, "
         "tail) of the input. All prior embedding approaches truncate to a prefix, "
         "which fails at long context lengths."),
        ("Position-Invariant Text Encoding for Cache Matching (Claims 4, 12, 17(d), 20)",
         "No prior system sorts text segments prior to embedding to achieve "
         "reorder-invariant semantic cache matching."),
        ("Tenant Affinity in Semantic Cache Routing (Claims 5, 14)",
         "No prior system applies per-tenant scoring bonuses in semantic KV cache "
         "routing based on empirical intra-tenant overlap analysis."),
    ]

    for title, desc in differentiators:
        numbered_para(f"{title}: {desc}", bold_first_sentence=True)

    # --- PRODUCTION EVIDENCE ---
    section_heading("PRODUCTION EVIDENCE AND REDUCTION TO PRACTICE")

    section_heading("Implementation Status", level=2)

    numbered_para(
        'The system has been implemented in production-quality code and deployed '
        'on GPU infrastructure with a plurality of GPU workers, a fleet router, '
        'and a distributed message bus for event-driven cache state synchronization. '
        'The implementation includes the semantic fleet router, fleet-level vector '
        'index with auto-adaptive backend, cascade discovery pipeline, multi-region '
        'prompt sampling, position-invariant encoding, and tenant affinity scoring. '
        'Partial KV reuse on workers uses known techniques (rotary position '
        'correction, selective recomputation) with the addition of per-architecture '
        'divergence profiling.'
    )

    section_heading("Empirical Validation", level=2)

    numbered_para(
        'The system has been validated across multiple model architectures, '
        'multiple NLP datasets (news summarization, dialogue, how-to articles, '
        'multi-document synthesis), and real-world conversational data comprising '
        'tens of thousands of query pairs from thousands of users. Key validated '
        'findings:'
    )

    findings = [
        'Fleet semantic routing successfully directs queries to workers holding '
        'relevant cached states, achieving semantic cache hit rates substantially '
        'higher than exact-prefix-only routing.',
        'Cascade discovery preserves existing exact-match performance while adding '
        'semantic coverage for queries that would otherwise go to cold fallback.',
        'Multi-region prompt sampling restores full cache hit rates on long-context '
        'inputs where prefix-only truncation produces zero hits.',
        'Tenant affinity scoring improves routing precision in multi-tenant '
        'workloads, consistent with the empirical finding that same-user query '
        'pairs exhibit substantially higher semantic overlap than cross-user pairs.',
        'The system correctly rejects dissimilar workloads (validated on code '
        'generation as negative control) with minimal overhead, confirming that '
        'the strictly additive architecture does not degrade unrelated workloads.',
    ]

    for finding in findings:
        numbered_para(finding)

    # --- FILING STRATEGY ---
    section_heading("FILING STRATEGY")

    numbered_para(
        'File provisional application immediately to secure priority date for '
        'semantic fleet routing, cascade discovery, multi-region prompt sampling, '
        'and position-invariant encoding claims.'
    )

    numbered_para(
        'File non-provisional CIP within 12 months, claiming priority to Filing 1 '
        'and Filing 3 for commonly disclosed matter and to the provisional for '
        'new matter.'
    )

    numbered_para(
        'Prosecution budget estimated for 2 Office Action responses.'
    )

    numbered_para(
        'Priority Date Mapping: Claims directed to similarity threshold mechanism '
        'and embedding-based matching: priority of Filing 1. All other claims '
        '(fleet routing with cache affinity, cascade discovery, event-driven '
        'synchronization, multi-region sampling, position-invariant encoding, '
        'tenant affinity scoring): priority of this application.'
    )

    # --- SUBJECT MATTER ELIGIBILITY ---
    section_heading("SUBJECT MATTER ELIGIBILITY DECLARATION")

    numbered_para(
        'Under the USPTO\u2019s updated guidance (Ex Parte Desjardins, September 2025; '
        'Kim Memo, August 2025), all claims are patent-eligible under Alice. Claims '
        'recite concrete computational operations: maintaining and querying a '
        'fleet-level vector index (Claims 1\u20132), performing cascade hash-then-semantic '
        'lookup (Claims 3, 9), extracting text from multiple regions of a token '
        'sequence (Claim 17), sorting text segments for position-invariant encoding '
        '(Claim 20), and routing inference queries based on computed similarity '
        'scores (Claim 9). These are specific, technical operations producing '
        'measurable performance improvements in GPU-based inference systems.'
    )

    # --- PROSECUTION STRATEGY ---
    section_heading("PROSECUTION STRATEGY")

    numbered_para(
        'If examiner combines existing KV-aware router with existing semantic cache '
        'with existing KV sharing: (1) existing KV-aware routers match on exact '
        'prefix only, not semantic similarity of cached KV states; (2) application-level '
        'semantic caches cache complete responses, not KV tensor states\u2014they cannot '
        'enable partial prefill reuse; (3) single-instance KV sharing systems '
        'operate on a single instance, not across a fleet\u2014they have no concept of '
        'fleet routing or cross-worker cache state synchronization; (4) no prior '
        'system combines exact hash and semantic search in a strictly additive '
        'cascade; (5) no prior system addresses the embedding truncation problem '
        'for long-context cache matching via multi-region sampling.'
    )

    numbered_para(
        'If examiner cites vLLM Semantic Router: vLLM Semantic Router routes '
        'between DIFFERENT MODELS (Mixture-of-Models routing) based on query '
        'semantics. The present invention routes between instances of the SAME '
        'MODEL based on each instance\u2019s cached KV tensor states. These are '
        'fundamentally different problems with different architectures.'
    )

    numbered_para(
        'If 35 U.S.C. \u00a7 101 rejection: cite Desjardins and Kim Memo. Claims '
        'recite specific technical operations on GPU infrastructure (vector index '
        'maintenance, hash computation, multi-region text extraction, routing '
        'decisions). Measured improvements demonstrate practical application.'
    )

    numbered_para(
        'If enablement challenge: production implementation constitutes constructive '
        'reduction to practice with validated empirical results across multiple '
        'model architectures, multiple NLP datasets, and real-world conversational data.'
    )

    # -- Add page numbers --
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Page number field
        run = p.add_run()
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)
        run2 = p.add_run()
        instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instr_text)
        run3 = p.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_char_end)

    # -- Save --
    doc.save(output_path)
    print(f"Patent document saved to: {output_path}")


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else (
        "/Users/zach/dev/worldflowai/autoresearch-semblend/patent/"
        "WorldFlowAI_Filing4_v5.0_Patent_Application.docx"
    )
    Path(out).parent.mkdir(parents=True, exist_ok=True)
    create_patent_document(out)
